"""
Generates a properly formatted El-Moquwal_Documentation.docx from the markdown source.
Run from the project root or the DOC folder:
    python DOC/generate_docx.py
"""
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent
MD_PATH = BASE / "El-Moquwal_Documentation.md"
REF_PATH = BASE / "reference.docx"
OUT_PATH = BASE / "El-Moquwal_Documentation.docx"
PREPROCESSED_PATH = BASE / "_preprocessed.md"  # temp file, deleted after run

NAVY = RGBColor(0x1F, 0x35, 0x64)
BLUE = RGBColor(0x2E, 0x74, 0xB5)

HEADING_STYLES = {f"Heading {i}" for i in range(1, 7)}
CODE_STYLES = {"Source Code", "Verbatim", "Code", "Verbatim Char"}
SKIP_JUSTIFY = HEADING_STYLES | CODE_STYLES | {"Caption", "Image Caption", "Figure Caption"}


def _pf(style):
    return style.paragraph_format


def create_reference():
    doc = Document()

    # A4 page, 2.54 cm margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Normal (body text)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = _pf(normal)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(18)  # 1.5× for 12pt

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    pf = _pf(h1)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.page_break_before = True
    pf.keep_with_next = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    pf = _pf(h2)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = NAVY
    pf = _pf(h3)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True

    # Heading 4
    try:
        h4 = doc.styles["Heading 4"]
        h4.font.name = "Calibri"
        h4.font.size = Pt(11)
        h4.font.bold = True
        h4.font.italic = True
        pf = _pf(h4)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        pf.keep_with_next = True
    except Exception:
        pass

    # Body Text
    try:
        bt = doc.styles["Body Text"]
        bt.font.name = "Times New Roman"
        bt.font.size = Pt(12)
        pf = _pf(bt)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(6)
        pf.line_spacing = Pt(18)
    except Exception:
        pass

    # Block Text (blockquotes)
    try:
        bq = doc.styles["Block Text"]
        bq.font.name = "Times New Roman"
        bq.font.size = Pt(11)
        bq.font.italic = True
        pf = _pf(bq)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Cm(1.27)
        pf.right_indent = Cm(1.27)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
    except Exception:
        pass

    # Source Code / Verbatim
    try:
        sc = doc.styles["Source Code"]
        sc.font.name = "Courier New"
        sc.font.size = Pt(9)
        pf = _pf(sc)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.left_indent = Cm(0.63)
    except Exception:
        pass

    try:
        vc = doc.styles["Verbatim Char"]
        vc.font.name = "Courier New"
        vc.font.size = Pt(9)
    except Exception:
        pass

    # Caption styles
    for cap_name in ("Caption", "Image Caption", "Figure Caption"):
        try:
            cap = doc.styles[cap_name]
            cap.font.name = "Times New Roman"
            cap.font.size = Pt(10)
            cap.font.italic = True
            _pf(cap).alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Placeholder paragraph required by Pandoc
    doc.add_paragraph("Reference document — do not edit.", style="Normal")

    doc.save(str(REF_PATH))
    print(f"[1/3] Reference doc created: {REF_PATH}")


def preprocess_markdown():
    """Fix image paths: decode %20 and convert file:/// URLs to relative paths."""
    text = MD_PATH.read_text(encoding="utf-8")

    def fix_img(m):
        alt = m.group(1)
        url = m.group(2)
        # Decode URL-encoded chars and strip the file:/// scheme
        if url.startswith("file:///"):
            decoded = unquote(url[8:])  # "C:/Projects/el-moquwal/UI/Screenshot 2026-..."
            p = Path(decoded)
            if p.exists():
                url = p.as_posix()  # absolute POSIX path — Pandoc handles these reliably
            else:
                # Try locating via UI folder
                fname = p.name
                candidate = BASE.parent / "UI" / fname
                if candidate.exists():
                    url = candidate.as_posix()
        return f"![{alt}]({url})"

    processed = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", fix_img, text)
    PREPROCESSED_PATH.write_text(processed, encoding="utf-8")
    print(f"[1.5/3] Markdown preprocessed (image paths fixed)")
    return PREPROCESSED_PATH


def run_pandoc():
    src = preprocess_markdown()
    cmd = [
        "pandoc",
        str(src),
        "-o", str(OUT_PATH),
        f"--reference-doc={REF_PATH}",
        "--from", "markdown+raw_tex-yaml_metadata_block",
        "--columns=10000",
    ]
    print(f"[2/3] Running Pandoc...")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.stdout.strip():
        print(result.stdout)
    if result.stderr.strip():
        print(result.stderr, file=sys.stderr)
    # Clean up temp file
    try:
        PREPROCESSED_PATH.unlink()
    except Exception:
        pass
    if result.returncode != 0:
        print("Pandoc failed — see error above.", file=sys.stderr)
        sys.exit(result.returncode)
    print(f"      Output: {OUT_PATH}")


def post_process():
    doc = Document(str(OUT_PATH))

    for para in doc.paragraphs:
        name = para.style.name
        if not para.text.strip():
            continue
        if name in HEADING_STYLES:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif name not in SKIP_JUSTIFY:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if para.paragraph_format.line_spacing is None:
                para.paragraph_format.line_spacing = Pt(18)

    doc.save(str(OUT_PATH))
    print(f"[3/3] Post-processing done.")


if __name__ == "__main__":
    create_reference()
    run_pandoc()
    post_process()
    print(f"\nDone! Open: {OUT_PATH}")
